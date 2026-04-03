from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pdfplumber
from docx import Document

from .validators import AppError


class FileHandler:
    def extract_text(self, filename: str, content: bytes) -> str:
        extension = Path(filename).suffix.lower()

        if extension == ".pdf":
            return self._extract_pdf(content)

        if extension == ".docx":
            return self._extract_docx(content)

        if extension in {".txt", ".text"}:
            return self._extract_text(content)

        raise AppError(
            status_code=415,
            code="unsupported_file_type",
            message="Only PDF, DOCX, and TXT files are supported.",
        )

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        pages: list[str] = []

        with pdfplumber.open(BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=1.5, y_tolerance=2) or ""
                if text.strip():
                    pages.append(text.strip())

        return "\n\n".join(pages)

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        document = Document(BytesIO(content))
        paragraphs = [item.text.strip() for item in document.paragraphs if item.text.strip()]
        table_rows: list[str] = []

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_rows.append(row_text)

        return "\n".join(paragraphs + table_rows)

    @staticmethod
    def _extract_text(content: bytes) -> str:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1", errors="ignore")
